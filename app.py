import os 
import pandas as pd 
import streamlit as st 
import io 
import urllib.parse 
from datetime import datetime, timedelta, timezone 
from docx import Document 

st.set_page_config(page_title="Gerador de Memorandos", page_icon="📄", layout="wide") 
st.title("📄 Emissor de Memorandos Individuais - Hospital Dr. Jackson Lago") 
st.markdown("### 🗓️ Configuração da Data de Envio") 

data_selecionada = st.date_input("Selecione a data que sairá no cabeçalho do Memorando:", datetime.now()) 
arquivo_excel = st.file_uploader("Suba a planilha contendo os incidentes (.xlsx)", type=["xlsx"]) 
caminho_modelo = "modelo_memorando.docx" 

def substituir_texto_protegendo_logos(doc, dicionario_tags): 
    """Substitui o texto alterando apenas os 'runs' para proteger imagens e cabeçalhos.""" 
    for paragrafo in doc.paragraphs: 
        for tag, valor in dicionario_tags.items(): 
            if tag in paragrafo.text: 
                for run in paragrafo.runs: 
                    if tag in run.text: 
                        run.text = run.text.replace(tag, valor) 
    for tabela in doc.tables: 
        for linha in tabela.rows: 
            for celula in linha.cells: 
                for paragrafo in celula.paragraphs: 
                    for tag, valor in dicionario_tags.items(): 
                        if tag in paragrafo.text: 
                            for run in paragrafo.runs: 
                                if tag in run.text: 
                                    run.text = run.text.replace(tag, valor) 

def formatar_data_br(valor): 
    """Garante que as datas sejam exibidas no formato brasileiro DD/MM/AAAA""" 
    try: 
        if pd.isna(valor) or str(valor).strip() == "" or str(valor).strip().lower() == "nan": 
            return "" 
        return pd.to_datetime(valor).strftime("%d/%m/%Y") 
    except: 
        return str(valor).strip() 

def limpar_numero_float(valor): 
    """Remove o .0 de números inteiros vindos do Excel (ex: 886.0 vira 886)""" 
    if pd.isna(valor) or str(valor).strip().lower() == "nan": 
        return "" 
    try: 
        if isinstance(valor, float) and valor.is_integer(): 
            return str(int(valor)) 
        v_str = str(valor).strip() 
        if v_str.endswith(".0"): 
            return v_str[:-2] 
        return v_str 
    except: 
        return str(valor) 

def tratar_str_limpa(valor): 
    """Evita que campos vazios ou nulos exibam a palavra 'nan'""" 
    if pd.isna(valor) or str(valor).strip().lower() == "nan" or str(valor).strip().lower() == "none": 
        return "" 
    return str(valor).strip() 

def obter_data_por_extenso(dt): 
    """Gera a data selecionada por extenso em português brasileiro""" 
    meses = { 
        1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho", 
        7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro" 
    } 
    return f"{dt.day} de {meses[dt.month]} de {dt.year}" 

@st.fragment 
def renderizar_bloco_memorando(index, id_bloco, dados_gerais, dados_especificos, data_extenso_envio):
    """Renderiza um card específico de memorando na interface para o usuário."""
    num_notif = dados_gerais["num_notif"]
    nome_do_paciente = dados_gerais["nome_do_paciente"]
    
    num_memo_cru = dados_especificos["num_memo"]
    gestor_destinatario = dados_especificos["gestor"]
    setor_notificado_final = dados_especificos["setor"]
    email_destino = dados_especificos["email"]

    num_memo_limpo = num_memo_cru.replace("Nº", "").replace("NS", "").replace("NSP", "").replace("/", "-").replace(" ", "").strip() 
    nome_base_arquivo = f"MEMORANDO Nº {num_memo_limpo}_NOTIFICAÇÃO_Nº {num_notif}_I_NSP" 
    
    dados_memorando = { 
        "{{numero_memorando}}": num_memo_cru, 
        "{{gestor}}": gestor_destinatario, 
        "{{setor}}": setor_notificado_final, 
        "{{notificacao_n}}": num_notif, 
        "{{data_notificacao}}": dados_gerais["dt_notif"], 
        "{{data_ocorrencia}}": dados_gerais["dt_ocorr"], 
        "{{localizacao}}": dados_gerais["onde_ocorreu"], 
        "{{tipo_incidente}}": dados_gerais["tipo_incidente"], 
        "{{classificacao_incidente}}": dados_gerais["classificacao_incidente"], 
        "{{descricao_notificacao}}": dados_gerais["descricao_notificacao"], 
        "{{nome_paciente}}": nome_do_paciente, 
        "{{leito}}": dados_gerais["leito_paciente"], 
        "{{setor_notificante}}": dados_gerais["setor_notificante_bruto"], 
        "{{sugestao}}": dados_gerais["sugestao_nsp"], 
        "{{m}}": dados_gerais["marca_manha"], 
        "{{t}}": dados_gerais["marca_tarde"], 
        "{{n}}": dados_gerais["marca_noite"], 
        "{{data_envio}}": data_extenso_envio 
    } 
    
    fuso_brasilia = timezone(timedelta(hours=-3)) 
    hora_atual = datetime.now(fuso_brasilia).hour 
    saudacao = "Bom Dia Prezados" if hora_atual < 12 else "Boa Tarde Prezados" 
    
    corpo_email = ( 
        f"{saudacao},\n\n" 
        f"Segue em Anexo o Memorando Nº {num_memo_cru} para ser analisado e respondido " 
        f"(via e-mail) em até 15 dias após a data presente.\n\n" 
        f"ATENÇÃO: A resposta via e-mail deve constar um arquivo em forma de word ou PDF para " 
        f"arquivamento de respostas conforme rotina institucional. Não serão aceitas mensagens " 
        f"via e-mail sem arquivo como resposta.\n\n" 
        f"Segue abaixo a notificação para análise do incidente em equipe e resposta ao NSP:\n" 
        f"• Memorando: Nº {num_memo_cru}\n" 
        f"• Notificação: Nº {num_notif}\n\n" 
        f"Atenciosamente,\n" 
        f"Ezequias S. Santos\n" 
        f"Agente Administrativo" 
    ) 
    
    col_nome, col_word, col_pdf, col_copiar = st.columns([1.5, 0.8, 0.8, 1.8]) 
    
    with col_nome: 
        st.markdown(f"**🔹 {nome_do_paciente}**") 
        st.caption(f"📍 Destino: **{gestor_destinatario}** ({setor_notificado_final})")
        if email_destino: 
            st.caption(f"📧 Destinatário: {email_destino}") 
            
    with col_word: 
        doc_word = Document(caminho_modelo) 
        substituir_texto_protegendo_logos(doc_word, dados_memorando) 
        word_io = io.BytesIO() 
        doc_word.save(word_io) 
        word_io.seek(0) 
        st.download_button( 
            label="📝 WORD", 
            data=word_io.getvalue(), 
            file_name=f"{nome_base_arquivo}.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            key=f"w_{index}_{id_bloco}" 
        ) 
        
    with col_pdf: 
        st.download_button( 
            label="📕 PDF", 
            data=word_io.getvalue(), 
            file_name=f"{nome_base_arquivo}.pdf", 
            mime="application/pdf", 
            key=f"p_{index}_{id_bloco}" 
        ) 
        
    with col_copiar: 
        st.code(corpo_email, language="text") 
        
    st.markdown("---") 

def processar_linha_paciente_sob_demanda(index, linha, num_colunas, data_extenso_envio):
    # Criação de um dicionário nomeado para blindar os dados e evitar cortes da inteligência artificial
    info = {}
    for i, val in enumerate(linha):
        info[int(i)] = val

    num_notif = limpar_numero_float(info.get(0, "")) if num_colunas > 0 else "S-N" 
    if num_notif.upper() == "STATUS" or "NOTIF" in num_notif.upper() or num_notif == "" or num_notif == "1": 
        return 
        
    nome_do_paciente = tratar_str_limpa(info.get(9, "")) if num_colunas > 9 else "Paciente Não Identificado" 
    if nome_do_paciente.upper() == "PACIENTE": 
        return 

    dt_ocorr = formatar_data_br(info.get(2, "")) if num_colunas > 2 else "" 
    dt_notif = formatar_data_br(info.get(3, "")) if num_colunas > 3 else "" 
    turno_planilha = str(info.get(4, "")).strip().upper() if num_colunas > 4 else "" 
    onde_ocorreu = tratar_str_limpa(info.get(5, "")) if num_colunas > 5 else "Ala B" 
    tipo_incidente = tratar_str_limpa(info.get(6, "")) if num_colunas > 6 else "" 
    classificacao_incidente = tratar_str_limpa(info.get(7, "")) if num_colunas > 7 else "" 
    descricao_notificacao = tratar_str_limpa(info.get(8, "")) if num_colunas > 8 else "" 
    leito_paciente = limpar_numero_float(info.get(10, "")) if num_colunas > 10 else "" 
    setor_notificante_bruto = tratar_str_limpa(info.get(11, "")) if num_colunas > 11 else "" 
    sugestao_nsp = tratar_str_limpa(info.get(12, "")) if num_colunas > 12 else "" 
    
    if setor_notificante_bruto == "": 
        setor_notificante_bruto = "NSP - NÚCLEO DE SEGURANÇA DO PACIENTE" 

    marca_manha = "X" if "MANH" in turno_planilha else " " 
    marca_tarde = "X" if "TARD" in turno_planilha else " " 
    marca_noite = "X" if "NOIT" in turno_planilha else " " 

    dados_gerais = {
        "num_notif": num_notif, "nome_do_paciente": nome_do_paciente, "dt_ocorr": dt_ocorr,
        "dt_notif": dt_notif, "onde_ocorreu": onde_ocorreu, "tipo_incidente": tipo_incidente,
        "classificacao_incidente": classificacao_incidente, "descricao_notificacao": descricao_notificacao,
        "leito_paciente": leito_paciente, "setor_notificante_bruto": setor_notificante_bruto,
        "sugestao_nsp": sugestao_nsp, "marca_manha": marca_manha, "marca_tarde": marca_tarde, "marca_noite": marca_noite
    }
    
    blocos_destinos = []
    
    # Bloco 1 (Original: N=13, O=14, P=15, V=21)
    num_memo_1 = tratar_str_limpa(info.get(15, "")) if num_colunas > 15 else ""
    if num_memo_1 == "" or num_memo_1.upper() == "Nº MEMO 01": num_memo_1 = "S-N"
    email_1 = tratar_str_limpa(info.get(21, "")) if num_colunas > 21 else ""
    if email_1.upper() == "EMAIL_SETOR": email_1 = ""
    
    blocos_destinos.append({
        "setor": tratar_str_limpa(info.get(14, "")) if num_colunas > 14 else onde_ocorreu,
        "gestor": tratar_str_limpa(info.get(13, "")) if num_colunas > 13 else "GESTOR DE ENFERMAGEM",
        "num_memo": num_memo_1,
        "email": email_1
    })
    
